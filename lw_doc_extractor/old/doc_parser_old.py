"""
"""

import logging

import re
import collections
import json
import os
import errno

ID_MATCHER = re.compile("[A-Za-z\\-0-9]+")

logger = logging.getLogger(__name__)

# if __name__ == "__main__":
#     print(ID_MATCHER.match("G-Staircase-002"))


def check_id(id):
    if ID_MATCHER.match(id):
        return True
    else:
        return False


CHARATERS = ["BEN", "ADA", "GABRIEL", "ZEALOT 1", "ZEALOT 2", "AARON", "MAN 1", "MAN 2", "MAN 3"]

FLOW_FRAGMENTS = ["A-Booth-001", "A-Booth-002", "B-Marketplace-001"]

OPTIONS_STARTSTR = "Player options"

REFENCE_MATCHER = re.compile("§\s*([A-Za-z0-9\-_]+)")

import docx

def export_imgs(inputWordDoc, outputImgPath):
    # for s in doc.inline_shapes:
    #     logger.info(f"shape {s}: {s.width}x{s.height} {s.type}")
    #
    #
    #     inline = s._inline
    #     print("==============")
    #     print(inline)
    #     print("==============")
    #     rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
    #     print(rId)
    #     image_part = doc.part.related_parts[rId]
    #     image_bytes = image_part.blob
    #
    #     print(image_bytes)
    
    
    logger.info("Starting export of images")
    import aspose.words as aw
    # load the Word document
    doc = aw.Document(inputWordDoc)
    # retrieve all shapes
    shapes = doc.get_child_nodes(aw.NodeType.SHAPE, True)
    # loop through shapes
    i = 0
    for shape in shapes:
        shape = shape.as_shape()
        if (shape.has_image):
            # set image file's name
            imageTypeExt = aw.FileFormatUtil.image_type_to_extension(shape.image_data.image_type)
            imageFileNm = os.path.join(outputImgPath, f"Image{i}{imageTypeExt}")
            shape.image_data.save(imageFileNm)
            logger.debug(f"Exported Image {imageFileNm}")
            i += 1
            #print(dir(shape.image_data))
            
            #logger.info(f"{shape.name}/{shape.title}")
            #print(shape.source_full_name)
            # 
    logger.info(f"Exporting {i} images complete")
    
    return


def create_nodedict_from_node_content(nodeId, nodeContent):
    isDialog = nodeId in FLOW_FRAGMENTS
    nodeResDict = {"id": nodeId, "type": "fragment" if isDialog else "dialog"}
    inPlayerOptions = False
    internal_content = []
    internal_links = []
    external_links = []
    currIntNode = nodeId
    for i, aline in enumerate(nodeContent):
        if not aline.strip() or aline.startswith("Set"):
            continue
        
        startsWithCharacter = False
        parts = aline.split(":")
        if len(parts) > 1:
            for ch in CHARATERS:
                if inPlayerOptions:
                    charNm = parts[0].split(" ")[-1]
                else:
                    charNm = parts[0].strip()
                if charNm.startswith(ch):
                    startsWithCharacter = True
        if startsWithCharacter:
            intNodeId = nodeId + "_" + str(len(internal_content))
            line = parts[1].split("►")[0].strip()
            line = line.strip("‘").strip("’")
            internal_content.append({"type": "dialog_line",
                                     "character": charNm,
                                     "line": line,
                                     "id": intNodeId})
            internal_links.append((currIntNode, intNodeId))
            if not inPlayerOptions:
                currIntNode = intNodeId
            continue
        
        if aline.startswith("Player options"):
            inPlayerOptions = True
            continue
        
        if "►" in aline:
            strParts = aline.split("►")
            if len(strParts) != 2:
                logger.warning("Can't split link text: {p.text}")
            else:
                nStr = strParts[1].strip(")").strip()
                if check_id(currIntNode) and check_id(nStr):
                    logger.debug(f"Added new link '{currIntNode}': '{nStr}'")
                    external_links.append((currIntNode, nStr))
                else:
                    logger.warning(f"Link invalid syntax in line: {aline}")
            continue
        
        if i == 0:
            nodeResDict["description"] = aline
    
    nodeResDict["internal_content"] = internal_content
    nodeResDict["internal_links"] = internal_links
    nodeResDict["external_links"] = external_links
    return nodeResDict

def parse_node_definition_properties(augmented_lines):
    return {}

def split_node_definition(augmented_lines):
    
    propertySection = []
    startSection = []
    restRefSections = collections.OrderedDict()
    
    currentRefSection = None
    
    inPropertiesSection = True
    
    for augline in restRefSections[currentRefSection]:
        ref, line = augline
        if inPropertiesSection:
            if line.startswith("§"):
                raise RuntimeError(f"Cannot have an § reference before there is a single statement with - in {ref}")
            if line.startswith("-"):
                inPropertiesSection = False
                startSection.append(augline)
            else:
                propertySection.append(augline)
        else:
            m = REFENCE_MATCHER.match(line)
            if m:
                if currentRefSection:
                    if len(restRefSections[currentRefSection]) == 0:
                        raise RuntimeError(f"Section {currentRefSection} needs to have at least one statement '-'. {ref}")
                currentRefSection = m.group(1)
                restRefSections[currentRefSection] = []
            else:
                if not currentRefSection:
                    startSection.append(augline)
                else:
                    restRefSections[currentRefSection].append(augline)
                    
    return propertySection, startSection, restRefSections


def parse(inputWordDoc, outputJsonFilePath, outputImageDir, characterList):
    doc = docx.Document(inputWordDoc)
    logger.info(f"Doc number of paragraphs: {len(doc.paragraphs)}")
    logger.info(f"Doc number of inline shapes: {len(doc.inline_shapes)}")
    
    lines = []
    imageCounter = 0
    
    for i, p in enumerate(doc.paragraphs):
        #logger.debug(str(i)+" "+p.text)
        currrentLine = p.text.strip()
        if "<w:drawing>" in str(p._p.xml):
            imgIdStr = f"<IMAGE {imageCounter}>"
            logger.info(f"Inserted {imgIdStr}")
            lines.append(imgIdStr)
            imageCounter += 1
        if not currrentLine:
            continue
        lines.append(currrentLine)
        
    lines.append("") # add a trailing newline
        
    outputDirPath = os.path.dirname(outputJsonFilePath)
    debugDirPath = os.path.join(outputDirPath, "debug")
    
    try:
        os.makedirs(debugDirPath)
    except OSError as e:
        if e.errno != errno.EEXIST:
            raise

    rawOutputLocaiton = os.path.join(debugDirPath,"doc_output.raw")
    with open(rawOutputLocaiton, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    
    linesOutputLocaiton = os.path.join(debugDirPath,"doc_output.json")
    with open(linesOutputLocaiton, "w") as fh:
        json.dump(lines, fh, indent=2)
    
    logger.info(f"Extracted {len(lines)} lines. A copy was written to {linesOutputLocaiton}")
    from lw_doc_extractor import lexer, story_compiler
    ast = lexer.parse(lines, debugDirPath)
    
    logger.info(f"Lexing complete")
    
    lexOutPath = os.path.join(debugDirPath,"lexer_output.json")
    
    with open(lexOutPath, "w") as fh:
        json.dump(ast, fh, indent=2)
    
    logger.info(f"Parsed and structured (lexing) output written to {lexOutPath}")
    
    resultJson = story_compiler.compile_story(ast)
    
    logger.info(f"Compilation complete")
    
    with open(outputJsonFilePath, "w") as fh:
        json.dump(resultJson, fh, indent=2)
        
    logger.info(f"Final (compiled) output written to {outputJsonFilePath}")
    

def _old_parse(inputWordDoc, outputJsonFilePath, outputImageDir, characterList):
    doc = docx.Document(inputWordDoc)
    logger.info(f"Doc number of paragraphs: {len(doc.paragraphs)}")
    logger.info(f"Doc number of inline shapes: {len(doc.inline_shapes)}")
    
    #export_imgs(inputWordDoc, outputImageDir)
    
    # Matches node definition line (Chapter, section & nodes)
    nodeDefnMatcher = re.compile("\s*\*\s*([\[\]A-Za-z\-]+)\s+([A-Za-z\-_0-9]+)(.*)")
    
    nodeType = collections.OrderedDict()
    nodesContent = {}
    currentNodeId = None
    
    imageDict = {}
    
    lines = []
    
    for i, p in enumerate(doc.paragraphs):
        
        logger.debug(p.text)
        currrentLine = p.text.strip()
        pf = p.paragraph_format
        #logger.debug("Style: algnment {}, left indent {}, first line {}".format(pf.alignment, pf.left_indent, pf.first_line_indent))
        #print([s.type for s in p.part.inline_shapes])
        # print(p.runs)
        # print(dir(p.runs[0]))

#        print(p._p.xml)

        if not currrentLine:
            continue
        
        m = nodeDefnMatcher.match(currrentLine)
        if m:
            nodeId = m.group(2)
            nodesContent[nodeId] = []
            nodeType[nodeId] = m.group(1).strip('[').strip(']')
            currentNodeId = nodeId
            logger.info(f"id:{m.group(2)}, type:{nodeType[nodeId]}")
        elif currrentLine.startswith("*"):
            logger.warning(f"Line starts with * but was not detected as node definition: {currrentLine} at paragraph {i}")
            if currentNodeId:
                nodesContent.pop(currentNodeId)
                nodeType.pop(currentNodeId)
                currentNodeId = None
        else:
            # print( p._p.xml)
            # print("<w:drawing>" in str(p._p.xml))
            if "<w:drawing>" in str(p._p.xml):
                currImgId = len(imageDict)
                if not nodeId or nodeId in imageDict:
                    logger.warning("Image found, but nodeId either not set or already in imageDict")
                else:
                    imageDict[nodeId] = currImgId
                    logger.debug(f"Image for {nodeId}: {currImgId}")
                    lines.append(f"<img>image{currImgId}")
                    nodesContent[currentNodeId].append(f"<img>image{currImgId}")
            if currentNodeId:
                nodesContent[currentNodeId].append((i, currrentLine))
            
        lines.append(currrentLine)
        
    from lw_doc_extractor import lexer
    lexer.parse(lines)
                
    # for nodeId in nodesContent:
    #     nodePropertyLines, nodeMainRunLines, referenceToSubRunsDict = split_node_definition(nodesContent[currentNodeId])
    #
    #     nodePropertyDict = parse_node_definition_properties(nodePropertyLines)
    
    #print(json.dumps(nodesContent, indent=2))
    
        # if(p.text.startswith("*")):
    #         if(p.text.startswith("*E-Aaron-001")):
    #             logger.info(f"End found. Parsing done.")
    #             break
    #         nStr = p.text.strip("*").strip(")").strip()
    #         if check_id(nStr):
    #             logger.debug(f"Added new node '{nStr}'")
    #             currentNode = nStr
    #             nodes[currentNode] = []
    #             nodesContent[currentNode] = []
    #         else:
    #             logger.warning(f"New node invalid syntax in line: {p.text}")
    #
    #     elif currentNode:
    #         nodesContent[currentNode].append(p.text)
    #
    # print(json.dumps(nodesContent, indent=2))
    
    #
    #
    # finalDict = {"nodes" : []}
    # for nodeId in nodesContent:
    #     print("Processing {}".format(nodeId))
    #     nodeResDict = create_nodedict_from_node_content(nodeId, nodesContent[nodeId])
    #     print(json.dumps(nodeResDict, indent=2))
    #     finalDict["nodes"].append(nodeResDict)
    # #
    # # for node in nodes:
    # #     logger.info(f"'{node}':")
    # #     for l in nodes[node]:
    # #         logger.info(f" -> '{l}'")
    #
    # with open("out.json", "w") as fh:
    #     json.dump(finalDict, fh, indent=2)
    # logger.info("Done")



