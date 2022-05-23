"""
"""

import logging


import re
import collections

ID_MATCHER = re.compile("[A-Za-z\\-0-9]+")

logger = logging.getLogger(__name__)

# if __name__ == "__main__":
#     print(ID_MATCHER.match("G-Staircase-002"))


def check_id(id):
    if ID_MATCHER.match(id):
        return True
    else:
        return False


import docx
def do_something(postionalArg):
    
    nodes = collections.OrderedDict()
    doc = docx.Document(postionalArg)
    logger.info(f"Doc number of paragraphs: {len(doc.paragraphs)}")
    currentNode = None
    for p in doc.paragraphs:
        logger.debug(p.text)
        if(p.text.startswith("*")):
            nStr = p.text.strip("*").strip()
            if check_id(nStr):
                logger.debug(f"Added new node '{nStr}'")
                currentNode = nStr
                nodes[currentNode] = []
            else:
                logger.warning(f"New node invalid syntax in line: {p.text}")
            
        if "►" in p.text:
            strParts = p.text.split("►")
            if len(strParts) != 2:
                logger.warning("Can't split link text: {p.text}")
            else:
                nStr = strParts[1].strip()
                if check_id(currentNode) and check_id(nStr):
                    nodes[currentNode].append(nStr)
                    logger.debug(f"Added new link '{currentNode}': '{nStr}'")
                else:
                    logger.warning(f"Link invalid syntax in line: {p.text}")
        
    for node in nodes:
        logger.info(f"'{node}':")
        for l in nodes[node]:
            logger.info(f" -> '{l}'")

    logger.info("Done")