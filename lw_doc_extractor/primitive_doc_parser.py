"""
"""

import logging


import re
import collections
import json

ID_MATCHER = re.compile("[A-Za-z\\-0-9]+")

logger = logging.getLogger(__name__)

# if __name__ == "__main__":
#     print(ID_MATCHER.match("G-Staircase-002"))


def check_id(id):
    if ID_MATCHER.match(id):
        return True
    else:
        return False

CHARATERS = ["BEN", "ADA", "GABRIEL", "ZEALOT 1", "ZEALOT 2", "AARON", "MAN 1", "MAN 2", "MAN 3"]

FLOW_FRAGMENTS = ["A-Booth-001", "A-Booth-002", "B-Marketplace-001"]

OPTIONS_STARTSTR = "Player options"

import docx

def create_nodedict_from_node_content(nodeId, nodeContent):
    isDialog = nodeId in FLOW_FRAGMENTS
    nodeResDict = {"id": nodeId, "type": "fragment" if isDialog else "dialog"}
    inPlayerOptions = False
    internal_content = []
    internal_links = []
    external_links = []
    currIntNode = nodeId
    for i, aline in enumerate(nodeContent):
        if not aline.strip() or aline.startswith("Set"):
            continue
        
        startsWithCharacter = False
        parts = aline.split(":")
        if len(parts) > 1:
            for ch in CHARATERS:
                if inPlayerOptions:
                    charNm = parts[0].split(" ")[-1]
                else:
                    charNm = parts[0].strip()
                if charNm.startswith(ch):
                    startsWithCharacter = True
        if startsWithCharacter:
            intNodeId = nodeId + "_" + str(len(internal_content))
            line = parts[1].split("►")[0].strip()
            line = line.strip("‘").strip("’")
            internal_content.append({"type" : "dialog_line",
                                     "character": charNm,
                                     "line": line,
                                     "id": intNodeId})
            internal_links.append((currIntNode, intNodeId))
            if not inPlayerOptions:
                currIntNode = intNodeId
            continue
        
        if aline.startswith("Player options"):
            inPlayerOptions = True
            continue

        
        if "►" in aline:
            strParts = aline.split("►")
            if len(strParts) != 2:
                logger.warning("Can't split link text: {p.text}")
            else:
                nStr = strParts[1].strip(")").strip()
                if check_id(currIntNode) and check_id(nStr):
                    logger.debug(f"Added new link '{currIntNode}': '{nStr}'")
                    external_links.append((currIntNode, nStr))
                else:
                    logger.warning(f"Link invalid syntax in line: {aline}")
            continue
        
        if i == 0:
            nodeResDict["description"] = aline
    
    nodeResDict["internal_content"] = internal_content
    nodeResDict["internal_links"] = internal_links
    nodeResDict["external_links"] = external_links
    return nodeResDict

def parse(inputWordDoc, outputJsonFile):
    
    nodes = collections.OrderedDict()
    nodesContent = {}
    doc = docx.Document(inputWordDoc)
    logger.info(f"Doc number of paragraphs: {len(doc.paragraphs)}")
    currentNode = None
    for p in doc.paragraphs:
        logger.debug(p.text)
        pf = p.paragraph_format
        #logger.debug("Style: algnment {}, left indent {}, first line {}".format(pf.alignment, pf.left_indent, pf.first_line_indent))
        if(p.text.startswith("*")):
            if(p.text.startswith("*E-Aaron-001")):
                logger.info(f"End found. Parsing done.")
                break
            nStr = p.text.strip("*").strip(")").strip()
            if check_id(nStr):
                logger.debug(f"Added new node '{nStr}'")
                currentNode = nStr
                nodes[currentNode] = []
                nodesContent[currentNode] = []
            else:
                logger.warning(f"New node invalid syntax in line: {p.text}")
        
        elif currentNode:
            nodesContent[currentNode].append(p.text)

    print(json.dumps(nodesContent, indent=2))
    
    
    finalDict = {"nodes" : []}
    for nodeId in nodesContent:
        print("Processing {}".format(nodeId))
        nodeResDict = create_nodedict_from_node_content(nodeId, nodesContent[nodeId])
        print(json.dumps(nodeResDict, indent=2))
        finalDict["nodes"].append(nodeResDict)
    #
    # for node in nodes:
    #     logger.info(f"'{node}':")
    #     for l in nodes[node]:
    #         logger.info(f" -> '{l}'")

    with open("out.json", "w") as fh:
        json.dump(finalDict, fh, indent=2)
    logger.info("Done")