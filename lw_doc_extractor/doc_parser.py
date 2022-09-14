"""
"""

import logging

import re
import collections
import json
import os
import errno

ID_MATCHER = re.compile("[A-Za-z\\-0-9]+")

logger = logging.getLogger(__name__)

REFENCE_MATCHER = re.compile("§\s*([A-Za-z0-9\-_]+)")

import docx

def export_imgs(inputWordDoc, outputImgPath):
    # for s in doc.inline_shapes:
    #     logger.info(f"shape {s}: {s.width}x{s.height} {s.type}")
    #
    #
    #     inline = s._inline
    #     print("==============")
    #     print(inline)
    #     print("==============")
    #     rId = inline.xpath('./a:graphic/a:graphicData/pic:pic/pic:blipFill/a:blip/@r:embed')[0]
    #     print(rId)
    #     image_part = doc.part.related_parts[rId]
    #     image_bytes = image_part.blob
    #
    #     print(image_bytes)
    
    
    logger.info("Starting export of images")
    import aspose.words as aw
    # load the Word document
    doc = aw.Document(inputWordDoc)
    # retrieve all shapes
    shapes = doc.get_child_nodes(aw.NodeType.SHAPE, True)
    # loop through shapes
    i = 0
    for shape in shapes:
        shape = shape.as_shape()
        if (shape.has_image):
            # set image file's name
            imageTypeExt = aw.FileFormatUtil.image_type_to_extension(shape.image_data.image_type)
            imageFileNm = os.path.join(outputImgPath, f"Image{i}{imageTypeExt}")
            shape.image_data.save(imageFileNm)
            logger.debug(f"Exported Image {imageFileNm}")
            i += 1
            #print(dir(shape.image_data))
            
            #logger.info(f"{shape.name}/{shape.title}")
            #print(shape.source_full_name)
            # 
    logger.info(f"Exporting {i} images complete")
    
    return


def parse(inputWordDoc, outputImageDir, rawDebugOutputDir=None):
    doc = docx.Document(inputWordDoc)
    logger.info(f"Doc number of paragraphs: {len(doc.paragraphs)}")
    logger.info(f"Doc number of inline shapes: {len(doc.inline_shapes)}")
    
    lines = []
    imageCounter = 0
    
    for i, p in enumerate(doc.paragraphs):
        #logger.debug(str(i)+" "+p.text)
        currrentLine = p.text.strip()
        if "<w:drawing>" in str(p._p.xml):
            imgIdStr = f"<IMAGE {imageCounter}>"
            logger.info(f"Inserted {imgIdStr}")
            lines.append(imgIdStr)
            imageCounter += 1
        if not currrentLine:
            continue
        lines.append(currrentLine)
        
    lines.append("") # add a trailing newline
    
    logger.info(f"Extracted {len(lines)} lines.")
        
    if rawDebugOutputDir:
        rawOutputLocaiton = os.path.join(rawDebugOutputDir,"doc_output.raw")
        with open(rawOutputLocaiton, "w", encoding="utf-8") as fh:
            fh.write("\n".join(lines))
        
        linesOutputLocaiton = os.path.join(rawDebugOutputDir,"doc_output.json")
        with open(linesOutputLocaiton, "w") as fh:
            json.dump(lines, fh, indent=2)
    
            logger.info(f"A copy was written to {linesOutputLocaiton}")
    from lw_doc_extractor import lexer, story_compiler
    ast = lexer.parse(lines)
    
    logger.info(f"Lexing complete")
    
    return ast
    